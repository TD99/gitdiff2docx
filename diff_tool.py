import os
import json
import subprocess
import io
import mimetypes

import pathspec

from datetime import datetime

from PIL import Image

from difflib import SequenceMatcher

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from markdown_it.rules_block import paragraph

from pygments import lex
from pygments.lexers import guess_lexer_for_filename, guess_lexer
from pygments.styles import get_style_by_name

# The directory where this script is located
script_dir = os.path.dirname(os.path.abspath(__file__))

# ------------------------------------------------------------------------------
# Utility functions

def ask_yes_no(prompt: str, lang: dict) -> bool:
    while True:
        answer = input(prompt + " ").strip().lower()
        if answer in [lang["yes"], lang["no"]]:
            return answer == lang["yes"]

def print_green(text: str):
    print(f"\033[92m{text}\033[0m")

def print_yellow(text: str):
    print(f"\033[93m{text}\033[0m")

def print_red(text: str):
    print(f"\033[91m{text}\033[0m")

def is_binary_string(bytes_data):
    textchars = bytearray({7, 8, 9, 10, 12, 13, 27}
                          | set(range(0x20, 0x100)) - {0x7f})
    return bool(bytes_data.translate(None, textchars))

def is_image_file(filename):
    mimetype, _ = mimetypes.guess_type(filename)
    return mimetype and mimetype.startswith("image/")

def get_usable_width(document):
    section = document.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    return page_width - left_margin - right_margin  # in EMUs

def remove_cell_border(cell, borders=("top", "left", "bottom", "right")):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders {}/>'.format(nsdecls('w')))
    for border in borders:
        border_element = parse_xml(
            f'<w:{border} w:val="nil" {nsdecls("w")}/>'
        )
        tcBorders.append(border_element)
    tcPr.append(tcBorders)

# ------------------------------------------------------------------------------
# Configuration

config_file = os.path.join(script_dir, "config.json")
if not os.path.exists(config_file):
    print_red(f"Error: Configuration file not found: {config_file}")
    exit()

with open(config_file, "r", encoding="utf-8") as f:
    config = json.load(f)

file_encoding = config.get("file_encoding", "utf-8")

# ------------------------------------------------------------------------------
# Pygments style configuration

pygments_style = config.get("pygments_style", "default")
try:
    pygments_style_obj = get_style_by_name(pygments_style)
    token_styles = pygments_style_obj.styles
except Exception:
    token_styles = get_style_by_name("default").styles

# ------------------------------------------------------------------------------
# Localization

lang_dir = os.path.join(script_dir, "lang")
if not os.path.exists(lang_dir):
    print_red(f"Error: Language directory not found: {lang_dir}")
    exit()

lang_choice = config.get("language", "en")
lang_file = os.path.join(lang_dir, f"{lang_choice}.json")
if not os.path.exists(lang_file):
    print_red(f"Language '{lang_choice}' not found, defaulting to English if available.")
    lang_file = os.path.join(lang_dir, "en.json")

with open(lang_file, "r", encoding="utf-8") as f:
    lang = json.load(f)

# ------------------------------------------------------------------------------
# Show banner

print_green(lang["title"])

# ------------------------------------------------------------------------------
# Input prompts

while True:
    target_dir = input(lang["enter_target_dir"] + " ").strip()
    if not os.path.isdir(target_dir):
        print_red(lang["invalid_target_dir"])
        continue
    if ".git" not in os.listdir(target_dir):
        print_red(lang["no_git_repo_found"].format(target_dir=target_dir))
        if ask_yes_no(lang["still_continue"], lang):
            break
        continue
    os.chdir(target_dir)
    break

# GDDIgnore
gdd_ignore_path = os.path.join(target_dir, config.get("gdd_ignore_file_name", ".gddignore"))
ignore_spec = None

if os.path.exists(gdd_ignore_path):
    with open(gdd_ignore_path, "r", encoding="utf-8") as f:
        ignore_spec = pathspec.PathSpec.from_lines("gitwildmatch", f)

# FIRST COMMIT
commit1 = input(lang["enter_commit1"] + " ").strip()
commit1_specified = bool(commit1)
if not commit1_specified:
    commit1 = subprocess.run(
        ["git", "rev-list", "--max-parents=0", "HEAD"],
        capture_output=True, text=True, encoding="utf-8"
    ).stdout.strip()

# Special case: if commit1 is the very first commit, we cannot add a caret
very_first_commit_hash = subprocess.run(
    ["git", "rev-list", "--max-parents=0", "HEAD"],
    capture_output=True, text=True, encoding="utf-8"
).stdout.strip()
is_very_first_commit = (commit1 == very_first_commit_hash)


if not (commit1.endswith("^") or "~" in commit1) and commit1 != "HEAD":
    include_first_commit = config.get("include_first_commit", False)

    if is_very_first_commit and include_first_commit:
        # Special revision number for an empty tree (state before any commit)
        empty_tree = "4b825dc642cb6eb9a060e54bf8d69288fbee4904"
        commit1 = empty_tree
    elif include_first_commit:
        commit1 = f"{commit1}^" if include_first_commit else commit1

if not commit1_specified:
    print(lang["using_first_commit"].format(commit1=commit1))

# LAST COMMIT
commit2 = input(lang["enter_commit2"] + " ").strip()
if not commit2:
    commit2 = subprocess.run(
        ["git", "rev-parse", "HEAD"],
        capture_output=True, text=True, encoding="utf-8"
    ).stdout.strip()
    print(lang["using_last_commit"].format(commit2=commit2))

output_docx = input(lang["enter_output_docx"] + " ").strip()
if not output_docx:
    output_docx = os.path.join(script_dir, "output.docx")
    print(lang["using_default_output"].format(output_docx=output_docx))

# ------------------------------------------------------------------------------
# Diff generation
changed_files = subprocess.run(
    ["git", "diff", "--name-only", commit1, commit2],
    capture_output=True, text=True, encoding="utf-8"
).stdout.splitlines()

if ignore_spec:
    changed_files = [f for f in changed_files if not ignore_spec.match_file(f)]

if not changed_files:
    print_yellow(lang["no_changes_found"].format(commit1=commit1, commit2=commit2))
    exit()

# ------------------------------------------------------------------------------
# Word document generation

doc = Document()
doc.add_heading(f"{lang['git_changes_report']} ({commit1} → {commit2})", level=1)
doc.add_paragraph(lang["report_generated_on"].format(
    date=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
))

# Add legend table
def add_legend_table(document):
    legend_table = document.add_table(rows=0, cols=2)
    legend_table.style = "Table Grid"

    legend_data = [
        (lang["legend_add"], config.get("add_color", "D0FFD0"), config.get("add_symbol", "+")),
        (lang["legend_remove"], config.get("remove_color", "FFD0D0"), config.get("remove_symbol", "-")),
        (lang["legend_neutral"], config.get("neutral_color", "F5F5F5"), config.get("neutral_symbol", "=")),
    ]

    for label, color, symbol in legend_data:
        column = legend_table.add_row().cells
        column[0].text = label

        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
        column[1]._element.get_or_add_tcPr().append(shading)
        column[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p = column[1].paragraphs[0]
        run = p.add_run(symbol)
        font = run.font
        font.name = config.get("diff_font", "Courier New")
        font.size = Pt(int(config.get("diff_font_size", 8)))

doc.add_heading(lang["legend"], level=config.get("heading_level", 2))
add_legend_table(doc)

doc.add_page_break()
doc.add_heading(lang["diffs"], level=config.get("heading_level", 2))

# Extract line numbers from git diff
def extract_line_numbers(diff_lines):
    line_numbers = []
    current_line = 0
    for line in diff_lines:
        if line.startswith("@@"):
            parts = line.split(" ")
            new_file_info = parts[2]
            start_line = int(new_file_info.split(",")[0][1:])
            current_line = start_line
        elif not line.startswith("-"):
            line_numbers.append(current_line)
            current_line += 1
    return line_numbers

# Add a formatted and syntax-highlighted code diff table
def add_diff_table(document, diff_lines, line_numbers, lexer):
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    table.columns[0].width = Cm(0.57)
    table.columns[1].width = get_usable_width(document) - Cm(0.57)

    add_symbol = config.get("add_symbol", "+")
    remove_symbol = config.get("remove_symbol", "-")
    neutral_symbol = config.get("neutral_symbol", "=")

    total_rows = len(diff_lines)

    for idx, (line, line_number) in enumerate(zip(diff_lines, line_numbers)):
        row_cells = table.add_row().cells
        symbol_cell = row_cells[0]
        code_cell = row_cells[1]

        # Apply background shading
        if line.startswith("+"):
            fill = config.get("add_color", "D0FFD0")
            symbol = add_symbol
        elif line.startswith("-"):
            fill = config.get("remove_color", "FFD0D0")
            symbol = remove_symbol
        else:
            fill = config.get("neutral_color", "F5F5F5")
            symbol = neutral_symbol

        # Background color for both cells
        for cell in (symbol_cell, code_cell):
            shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), fill))
            cell._element.get_or_add_tcPr().append(shading)

        # Determine which borders to remove
        borders_to_remove_symbol = ["right"]
        borders_to_remove_code = ["left"]

        # For inner rows, remove top and bottom borders
        if idx == 0:
            borders_to_remove_symbol.append("bottom")
            borders_to_remove_code.append("bottom")
        elif idx == total_rows - 1:
            borders_to_remove_symbol.append("top")
            borders_to_remove_code.append("top")
        else:
            borders_to_remove_symbol.extend(["top", "bottom"])
            borders_to_remove_code.extend(["top", "bottom"])

        # Hide the border
        remove_cell_border(symbol_cell, borders=borders_to_remove_symbol)
        remove_cell_border(code_cell, borders=borders_to_remove_code)

        # Set paragraph format for symbol cell
        symbol_paragraph = symbol_cell.paragraphs[0]
        symbol_paragraph.clear()
        run_sym = symbol_paragraph.add_run(symbol)
        run_sym.font.name = config.get("diff_font", "Courier New")
        run_sym.font.size = Pt(int(config.get("diff_font_size", 8)))

        # Clear and format code paragraph
        paragraph = code_cell.paragraphs[0]
        paragraph.clear()  # remove any auto-inserted text
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        # Strip symbol from content
        code_content = line[1:]

        # Lex and style each token
        for ttype, value in lex(code_content, lexer):
            value = value.rstrip('\n')
            if not value:
                value = "\u00A0" # Non-breaking space

            run = paragraph.add_run(value)
            run.font.name = config.get("diff_font", "Courier New")
            run.font.size = Pt(int(config.get("diff_font_size", 8)))

            style_str = token_styles.get(ttype)
            if style_str:
                for part in style_str.split():
                    if part == "bold":
                        run.bold = True
                    elif part == "italic":
                        run.italic = True
                    elif part.startswith("#") and len(part) == 7:
                        hexcode = part.lstrip("#")
                        try:
                            r = int(hexcode[0:2], 16)
                            g = int(hexcode[2:4], 16)
                            b = int(hexcode[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
                        except ValueError:
                            pass

# Add images to the document
def add_image(document, file_bytes, image_name):
    image_stream = io.BytesIO(file_bytes)
    try:
        image = Image.open(image_stream)
        width, _ = image.size
        image_stream.seek(0)  # rewind for docx
        max_width_inches = 6  # ~75% of page width (8 inches)
        width_inches = min(max_width_inches, width / image.info.get('dpi', (96, 96))[0])
        document.add_picture(image_stream, width=Inches(width_inches))
        document.paragraphs[-1].alignment = 1  # center
    except Exception as e:
        document.add_paragraph(lang["error_inserting_image"].format(image_name=image_name))

# ------------------------------------------------------------------------------
# Main loop

if os.path.exists(output_docx):
    if not ask_yes_no(lang["output_exists"].format(output_docx=output_docx), lang):
        print(lang["exiting"])
        exit()
    else:
        while True:
            try:
                with open(output_docx, "a", encoding="utf-8"):
                    break
            except Exception as e:
                print_red(lang["error_removing_file"].format(output_docx=output_docx, error=str(e)))
                input(lang["press_enter_to_retry"])

verbose = config.get("verbose", False)

for index, file in enumerate(changed_files):
    if not index == 0:
        doc.add_page_break()

    doc.add_heading(f"{lang['file']}: {file}", level=config.get("heading_level", 2) + 1)

    if verbose:
        print(lang["processing_file"].format(file=file))

    # Try to get binary contents
    try:
        old_bytes = subprocess.run(
            ["git", "show", f"{commit1}:{file}"],
            capture_output=True
        ).stdout
    except:
        old_bytes = b""

    try:
        new_bytes = subprocess.run(
            ["git", "show", f"{commit2}:{file}"],
            capture_output=True
        ).stdout
    except:
        new_bytes = b""

    # If binary but not image → skip
    if is_binary_string(new_bytes) or is_binary_string(old_bytes):
        if is_image_file(file):
            # Insert only if the image changed
            if old_bytes != new_bytes:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(lang['image_changed'])
                run.italic = True

                if config.get("include_images", True):
                    add_image(doc, new_bytes, file)
        else:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(lang['binary_file_skipped'])
            run.italic = True
        continue

    # Fallback for text diff
    old_content = old_bytes.decode(file_encoding, errors="ignore").splitlines()
    new_content = new_bytes.decode(file_encoding, errors="ignore").splitlines()

    # Choose lexer based on filename and content
    sample = "\n".join(new_content or old_content)
    try:
        lexer = guess_lexer_for_filename(file, sample)
    except:
        lexer = guess_lexer(sample or "")

    matcher = SequenceMatcher(None, old_content, new_content)
    diff_lines = []
    line_nums = []

    old_idx = new_idx = 0

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for line in new_content[j1:j2]:
                diff_lines.append(f" {line}")
                line_nums.append(new_idx + 1)
                new_idx += 1
                old_idx += 1
        elif tag == "replace":
            for line in old_content[i1:i2]:
                diff_lines.append(f"-{line}")
                line_nums.append(old_idx + 1)
                old_idx += 1
            for line in new_content[j1:j2]:
                diff_lines.append(f"+{line}")
                line_nums.append(new_idx + 1)
                new_idx += 1
        elif tag == "delete":
            for line in old_content[i1:i2]:
                diff_lines.append(f"-{line}")
                line_nums.append(old_idx + 1)
                old_idx += 1
        elif tag == "insert":
            for line in new_content[j1:j2]:
                diff_lines.append(f"+{line}")
                line_nums.append(new_idx + 1)
                new_idx += 1

    # Add Table if there are significant changes
    if diff_lines:
        add_diff_table(doc, diff_lines, line_nums, lexer)
    else:
        doc.add_paragraph(lang["no_significant_changes"], style="Italic")

    if verbose:
        print_green(lang["processing_done"].format(file=file))

try:
    doc.save(output_docx)
except Exception as e:
    print_red(lang["error_saving_file"].format(output_docx=output_docx, error=str(e)))
    exit()

print_green(lang["saving_report"].format(output_docx=output_docx))

if config.get("open_after_creation", False):
    try:
        os.startfile(output_docx)
    except Exception as e:
        print_red(lang["error_opening_file"].format(output_docx=output_docx, error=str(e)))
        exit()
